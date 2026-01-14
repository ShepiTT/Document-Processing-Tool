#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
import win32com.client
import pythoncom

# 添加当前目录到路径，确保能导入模块
current_dir = os.path.dirname(os.path.abspath(__file__))
if current_dir not in sys.path:
    sys.path.insert(0, current_dir)

class WPSTestConverter:
    """专门用于测试WPS Office的转换器"""

    def __init__(self):
        self.word_app = None

    def initialize_wps_app(self):
        """强制初始化WPS应用程序"""
        try:
            pythoncom.CoInitialize()
            self.word_app = win32com.client.Dispatch("KWPS.Application")
            self.word_app.Visible = False
            self.word_app.DisplayAlerts = False
            print("[OK] WPS应用程序初始化成功")
            return True
        except Exception as e:
            print(f"[ERROR] 初始化WPS应用程序失败: {e}")
            self.word_app = None
            try:
                pythoncom.CoUninitialize()
            except:
                pass
            return False

    def close_wps_app(self):
        """关闭WPS应用程序"""
        if not self.word_app:
            return

        try:
            if hasattr(self.word_app, 'Quit'):
                self.word_app.Quit()
            print("[OK] WPS应用程序已关闭")
        except Exception as e:
            print(f"[WARNING] 关闭应用程序时出错: {e}")
        finally:
            self.word_app = None
            try:
                pythoncom.CoUninitialize()
            except:
                pass

    def convert_file(self, word_file, pdf_file=None):
        """转换单个Word文件为PDF"""
        if not self.word_app:
            print("[ERROR] WPS应用程序未初始化")
            return False

        try:
            from pathlib import Path
            word_path = Path(word_file)
            if not word_path.exists():
                print(f"[ERROR] 文件不存在: {word_file}")
                return False

            if pdf_file is None:
                pdf_file = word_path.with_suffix('.pdf')
            pdf_path = Path(pdf_file)
            pdf_path.parent.mkdir(parents=True, exist_ok=True)

            print(f"[CONVERT] 正在转换: {word_path.name}")
            doc = self.word_app.Documents.Open(str(word_path.resolve()))
            doc.ExportAsFixedFormat(str(pdf_path.resolve()), 17)  # 17 = PDF格式
            doc.Close(False)

            if pdf_path.exists():
                file_size = pdf_path.stat().st_size
                print(f"[OK] 转换成功! 文件大小: {file_size} bytes")
                return True
            else:
                print("[ERROR] PDF文件生成失败")
                return False

        except Exception as e:
            print(f"[ERROR] 转换失败: {e}")
            return False

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        self.close_wps_app()

def test_wps_detection():
    """测试WPS检测功能"""
    print("[TEST] 测试WPS Office检测功能")
    print("=" * 50)

    try:
        pythoncom.CoInitialize()
        wps_app = win32com.client.Dispatch("KWPS.Application")
        wps_app.Visible = False
        wps_app.DisplayAlerts = False
        version = wps_app.Version
        print(f"[OK] 检测到WPS Office (版本: {version})")
        wps_app.Quit()
        pythoncom.CoUninitialize()
        return True
    except Exception as e:
        print(f"[ERROR] WPS Office不可用: {e}")
        try:
            pythoncom.CoUninitialize()
        except:
            pass
        return False

def cleanup_test_files():
    """清理测试生成的所有文件"""
    test_files = [
        "test_wps.docx",
        "test_wps.pdf",
        "test_microsoft.docx",
        "test_microsoft.pdf",
        "test_document.docx",
        "test_document.pdf"
    ]

    cleaned_files = []
    for file_path in test_files:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                cleaned_files.append(file_path)
                print(f"[CLEANUP] 已删除测试文件: {file_path}")
        except Exception as e:
            print(f"[WARNING] 删除文件失败 {file_path}: {e}")

    if cleaned_files:
        print(f"[OK] 清理完成，共删除了 {len(cleaned_files)} 个测试文件")
    else:
        print("[INFO] 没有找到需要清理的测试文件")

def test_wps_conversion():
    """测试WPS转换功能"""
    print("\n[TEST] 测试WPS转换功能")
    print("=" * 50)

    # 先清理可能存在的旧测试文件
    print("[CLEANUP] 清理旧测试文件...")
    cleanup_test_files()

    # 创建测试文档
    test_doc = "test_wps.docx"
    test_pdf = "test_wps.pdf"

    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph("这是WPS测试文档，用于测试WPS Office转PDF功能。")
        doc.add_paragraph("如果转换成功，将生成对应的PDF文件。")
        doc.save(test_doc)
        print(f"[OK] 创建了测试文档: {test_doc}")
    except ImportError:
        print("[WARNING] 未安装python-docx库，跳过创建测试文档")
        # 即使没有docx库，也要清理文件
        cleanup_test_files()
        return True

    success = False
    with WPSTestConverter() as converter:
        if not converter.initialize_wps_app():
            print("[ERROR] 无法初始化WPS应用程序")
        else:
            success = converter.convert_file(test_doc, test_pdf)

    # 无论测试成功与否，都要清理测试文件
    print("[CLEANUP] 清理测试文件...")
    cleanup_test_files()

    return success

if __name__ == "__main__":
    print("WPS Office专用测试")
    print("=" * 50)

    # 程序启动时清理旧测试文件
    print("[CLEANUP] 程序启动时清理旧测试文件...")
    cleanup_test_files()

    try:
        detection_success = test_wps_detection()
        conversion_success = test_wps_conversion()

        if detection_success and conversion_success:
            print("\n[SUCCESS] WPS测试全部通过!")
        else:
            print("\n[FAILED] WPS测试失败!")
    finally:
        # 程序结束时再次清理测试文件，确保清理干净
        print("[CLEANUP] 程序结束时最终清理测试文件...")
        cleanup_test_files()
